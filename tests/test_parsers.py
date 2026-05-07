"""测试所有文档解析器。"""

import os
import tempfile

import pytest


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield d


class TestPDFParser:
    def test_parse_pdf(self, tmp_dir):
        """创建一个最小 PDF 并解析。"""
        import fitz
        path = os.path.join(tmp_dir, "test.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Hello PDF World")
        doc.save(path)
        doc.close()

        from app.services.parsers.pdf_parser import parse_pdf
        result = parse_pdf(path)

        assert "Hello PDF World" in result.text
        assert len(result.pages) == 1
        assert result.pages[0]["page_num"] == 1


class TestDocxParser:
    def test_parse_docx(self, tmp_dir):
        from docx import Document
        path = os.path.join(tmp_dir, "test.docx")
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Some content here.")
        doc.save(path)

        from app.services.parsers.docx_parser import parse_docx
        result = parse_docx(path)

        assert "Title" in result.text
        assert "Some content here" in result.text
        assert len(result.sections) >= 1


class TestMarkdownParser:
    def test_parse_markdown(self, tmp_dir):
        path = os.path.join(tmp_dir, "test.md")
        with open(path, "w") as f:
            f.write("# Main Title\n\n## Section\n\nSome paragraph text.\n")

        from app.services.parsers.md_parser import parse_markdown
        result = parse_markdown(path)

        assert "Main Title" in result.text
        assert "Section" in result.text
        assert len(result.sections) == 2
        assert result.sections[0]["title"] == "Main Title"
        assert result.sections[0]["level"] == 1


class TestTxtParser:
    def test_parse_txt(self, tmp_dir):
        path = os.path.join(tmp_dir, "test.txt")
        with open(path, "w") as f:
            f.write("Line one\nLine two\n")

        from app.services.parsers.txt_parser import parse_txt
        result = parse_txt(path)

        assert "Line one" in result.text
        assert result.metadata["format"] == "text"


class TestCSVParser:
    def test_parse_csv(self, tmp_dir):
        path = os.path.join(tmp_dir, "test.csv")
        with open(path, "w") as f:
            f.write("Name,Age,City\nAlice,30,NYC\nBob,25,LA\n")

        from app.services.parsers.csv_parser import parse_csv
        result = parse_csv(path)

        assert "Name | Age | City" in result.text
        assert "Alice | 30 | NYC" in result.text
        assert result.metadata["row_count"] == 3


class TestExcelParser:
    def test_parse_xlsx(self, tmp_dir):
        import openpyxl
        path = os.path.join(tmp_dir, "test.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Col1", "Col2"])
        ws.append(["val1", "val2"])
        wb.save(path)
        wb.close()

        from app.services.parsers.excel_parser import parse_xlsx
        result = parse_xlsx(path)

        assert "Data" in result.text
        assert "Col1 | Col2" in result.text
        assert "val1 | val2" in result.text


class TestParserService:
    def test_dispatch_by_type(self, tmp_dir):
        path = os.path.join(tmp_dir, "test.txt")
        with open(path, "w") as f:
            f.write("hello")

        from app.services.parser_service import parse_document
        result = parse_document(path, "txt")
        assert "hello" in result.text

    def test_unknown_type(self, tmp_dir):
        from app.core.exceptions import ValidationException
        from app.services.parser_service import parse_document

        with pytest.raises(ValidationException, match="No parser"):
            parse_document("/tmp/x.xyz", "xyz")
