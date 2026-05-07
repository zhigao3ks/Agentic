"""DOCX 解析器。"""

from docx import Document as DocxDocument

from app.services.parsers.models import ParsedDocument


class DocxParser:
    @staticmethod
    def parse(file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        full_text_parts = []
        sections = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            full_text_parts.append(text)
            style = para.style.name if para.style else ""
            if "Heading" in style or "heading" in style or "标题" in style:
                sections.append({"title": text, "level": _heading_level(style)})

        # 提取表格
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            full_text_parts.append(table_text)

        return ParsedDocument(
            text="\n\n".join(full_text_parts),
            sections=sections,
            metadata={"paragraph_count": len(doc.paragraphs)},
        )


def _heading_level(style_name: str) -> int:
    for i in range(1, 7):
        if str(i) in style_name:
            return i
    return 1


def parse_docx(file_path: str) -> ParsedDocument:
    return DocxParser.parse(file_path)
